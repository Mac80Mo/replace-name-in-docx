import os
from docx import Document

def replace_name_in_docx(input_file, output_file, old_name, new_name):
    """ Öffnet ein Word-Dokument, ersetzt den Namen und speichert die bearbeitete Datei """
    doc = Document(input_file)
    
    # Durchlaufe alle Absätze und ersetze den Namen
    for para in doc.paragraphs:
        if old_name in para.text:
            para.text = para.text.replace(old_name, new_name)
    
    # Durchlaufe alle Tabellen (falls der Name in Tabellen vorkommt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_name in cell.text:
                    cell.text = cell.text.replace(old_name, new_name)
                    
    # Speichere das bearbeitete Dokument
    doc.save(output_file)
    print(f"Bearbeitet gespeichert als: {output_file}")

def process_multiple_files(source_dir, destination_dir, old_name, new_name, filename_suffix=""):
    """ Durchläuft alle .docx-Dateien im Quellordner und speichert sie mit optionalem Zusatz im Dateinamen """
    
    if not os.path.exists(source_dir):
        print(f"Ordner {source_dir} existiert nicht.")
        return
    
    # Zielordner erstellen, falls er nicht existiert
    if not os.path.exists(destination_dir):
        os.makedirs(destination_dir)

    for filename in os.listdir(source_dir):
        if filename.endswith(".docx"):
            input_file = os.path.join(source_dir, filename)

            # **Erhalte den Original-Dateinamen und füge den Zusatz an**
            name_part, ext = os.path.splitext(filename)
            new_filename = f"{name_part}_{filename_suffix}{ext}" if filename_suffix else filename  # Falls kein Zusatz, bleibt Originalname
            
            output_file = os.path.join(destination_dir, new_filename)

            replace_name_in_docx(input_file, output_file, old_name, new_name)

# Ordnerstruktur gemäß deinem Bild
source_dir = "./Source"  # Quellordner mit Originaldateien
destination_dir = "./Destination"  # Zielordner für bearbeitete Dateien
old_name = "Max Mustermann"  # Name, der ersetzt werden soll
new_name = "Neuer Name"  # Neuer Name
filename_suffix = ""  # **Hier den gewünschten Zusatz eingeben, kann auch leer sein ("")**

process_multiple_files(source_dir, destination_dir, old_name, new_name, filename_suffix)
